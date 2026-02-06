# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
import requests
import time
import gc
from datetime import datetime, timedelta, date
from operator import itemgetter
import json
import re
import base64
import io
import altair as alt
from supabase import create_client
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ImportaÃ§Ã£o condicional
try:
    from streamlit_javascript import st_javascript
except ImportError:
    st_javascript = None

# ImportaÃ§Ãµes de utilitÃ¡rios
from utils import (get_brazil_time, get_secret, send_to_chat)

# ============================================
# 1. CONFIGURAÃ‡Ã•ES E CONSTANTES (EQUIPE ID 1)
# ============================================
DB_APP_ID = 1        # ID da Fila desta equipe
LOGMEIN_DB_ID = 1    # ID do LogMeIn (COMPARTILHADO - SEMPRE 1)

CONSULTORES = sorted([
    "Alex Paulo", "Dirceu GonÃ§alves", "Douglas De Souza", "Farley Leandro", "Gleis Da Silva", 
    "Hugo Leonardo", "Igor Dayrell", "Jerry Marcos", "Jonatas Gomes", "Leandro Victor", 
    "Luiz Henrique", "Marcelo Dos Santos", "Marina Silva", "Marina Torres", "Vanessa Ligiane"
])

# Listas de OpÃ§Ãµes
REG_USUARIO_OPCOES = ["CartÃ³rio", "Gabinete", "Externo"]
REG_SISTEMA_OPCOES = ["Conveniados", "Outros", "Eproc", "Themis", "JPE", "SIAP"]
REG_CANAL_OPCOES = ["Presencial", "Telefone", "Email", "Whatsapp", "Outros"]
REG_DESFECHO_OPCOES = ["Resolvido - Cesupe", "Escalonado"]

OPCOES_ATIVIDADES_STATUS = ["HP", "E-mail", "WhatsApp PlantÃ£o", "HomologaÃ§Ã£o", "RedaÃ§Ã£o Documentos", "Outros"]

# URLs e Visuais
GIF_BASTAO_HOLDER = "https://media2.giphy.com/media/v1.Y2lkPTc5MGI3NjExa3Uwazd5cnNra2oxdDkydjZkcHdqcWN2cng0Y2N0cmNmN21vYXVzMiZlcD12MV9pbnRlcm5uYWxfZ2lmX2J5X2lkJmN0PWc/3rXs5J0hZkXwTZjuvM/giphy.gif"
GIF_LOGMEIN_TARGET = "https://media1.giphy.com/media/v1.Y2lkPTc5MGI3NjExZjFvczlzd3ExMWc2cWJrZ3EwNmplM285OGFqOHE1MXlzdnd4cndibiZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/mcsPU3SkKrYDdW3aAU/giphy.gif"
BASTAO_EMOJI = "ğŸ­" 
PUG2026_FILENAME = "Carnaval.gif" 
APP_URL_CLOUD = 'https://controle-bastao-equipe1.streamlit.app'

# Secrets
CHAT_WEBHOOK_BASTAO = get_secret("chat", "bastao")
WEBHOOK_STATE_DUMP = get_secret("webhook", "test_state")

# ============================================
# 2. OTIMIZAÃ‡ÃƒO E CONEXÃƒO
# ============================================

@st.cache_resource(ttl=3600)
def get_supabase():
    try: 
        return create_client(st.secrets["supabase"]["url"], st.secrets["supabase"]["key"])
    except Exception as e:
        st.cache_resource.clear()
        return None

@st.cache_data(ttl=3600, show_spinner=False)
def carregar_dados_grafico():
    sb = get_supabase()
    if not sb: return None, None
    try:
        res = sb.table("atendimentos_resumo").select("data").eq("id", DB_APP_ID).execute()
        if res.data:
            json_data = res.data[0]['data']
            if 'totais_por_relatorio' in json_data:
                df = pd.DataFrame(json_data['totais_por_relatorio'])
                return df, json_data.get('gerado_em', '-')
    except Exception as e:
        st.error(f"Erro grÃ¡fico: {e}")
    return None, None

@st.cache_data
def get_img_as_base64_cached(file_path):
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except: return None

# ============================================
# 3. REPOSITÃ“RIO
# ============================================

def clean_data_for_db(obj):
    if isinstance(obj, dict):
        return {k: clean_data_for_db(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_data_for_db(i) for i in obj]
    elif isinstance(obj, (datetime, date)):
        return obj.isoformat()
    elif isinstance(obj, timedelta):
        return obj.total_seconds()
    else:
        return obj

@st.cache_data(ttl=5, show_spinner=False)
def load_state_from_db():
    sb = get_supabase()
    if not sb: return {}
    try:
        response = sb.table("app_state").select("data").eq("id", DB_APP_ID).execute()
        if response.data and len(response.data) > 0:
            return response.data[0].get("data", {})
        return {}
    except Exception as e:
        return {}

def save_state_to_db(state_data):
    sb = get_supabase()
    if not sb: 
        st.error("Sem conexÃ£o para salvar.")
        return
    try:
        sanitized_data = clean_data_for_db(state_data)
        sb.table("app_state").upsert({"id": DB_APP_ID, "data": sanitized_data}).execute()
    except Exception as e:
        st.error(f"ğŸ”¥ ERRO DE ESCRITA NO BANCO: {e}")

# --- LOGMEIN DB ---
def get_logmein_status():
    sb = get_supabase()
    if not sb: return None, False
    try:
        res = sb.table("controle_logmein").select("*").eq("id", LOGMEIN_DB_ID).execute()
        if res.data:
            return res.data[0].get('consultor_atual'), res.data[0].get('em_uso', False)
    except: pass
    return None, False

def set_logmein_status(consultor, em_uso):
    sb = get_supabase()
    if not sb: return
    try:
        dados = {
            "consultor_atual": consultor if em_uso else None,
            "em_uso": em_uso,
            "data_inicio": datetime.now().isoformat()
        }
        sb.table("controle_logmein").update(dados).eq("id", LOGMEIN_DB_ID).execute()
    except Exception as e: st.error(f"Erro LogMeIn DB: {e}")

# ============================================
# 4. FUNÃ‡Ã•ES DE UTILIDADE E IP
# ============================================
def get_browser_id():
    if st_javascript is None: return "no_js_lib"
    js_code = """(function() {
        let id = localStorage.getItem("device_id");
        if (!id) {
            id = "id_" + Math.random().toString(36).substr(2, 9);
            localStorage.setItem("device_id", id);
        }
        return id;
    })();"""
    try:
        return st_javascript(js_code, key="browser_id_tag")
    except: return "unknown_device"

def get_remote_ip():
    try:
        from streamlit.web.server.websocket_headers import ClientWebSocketRequest
        ctx = st.runtime.scriptrunner.get_script_run_ctx()
        if ctx and ctx.session_id:
            session_info = st.runtime.get_instance().get_client(ctx.session_id)
            if session_info:
                request = session_info.request
                if isinstance(request, ClientWebSocketRequest):
                    if 'X-Forwarded-For' in request.headers:
                        return request.headers['X-Forwarded-For'].split(',')[0]
                    return request.remote_ip
    except: return "Unknown"
    return "Unknown"

# --- LIMPEZA DE MEMÃ“RIA ---
def memory_sweeper():
    if 'last_cleanup' not in st.session_state:
        st.session_state.last_cleanup = time.time()
        return
    if time.time() - st.session_state.last_cleanup > 300:
        st.session_state.word_buffer = None 
        gc.collect()
        st.session_state.last_cleanup = time.time()
    
    if 'last_hard_cleanup' not in st.session_state:
        st.session_state.last_hard_cleanup = time.time()
        
    if time.time() - st.session_state.last_hard_cleanup > 14400: # 4h
        st.cache_data.clear()
        gc.collect()
        st.session_state.last_hard_cleanup = time.time()

# --- LÃ“GICA DE FILA VISUAL ---
def get_ordered_visual_queue(queue, status_dict):
    if not queue: return []
    current_holder = next((c for c, s in status_dict.items() if 'BastÃ£o' in (s or '')), None)
    if not current_holder or current_holder not in queue: return list(queue)
    try:
        idx = queue.index(current_holder)
        return queue[idx:] + queue[:idx]
    except ValueError: return list(queue)

# ============================================
# 5. MANIPULAÃ‡ÃƒO DE ESTADO E DATA
# ============================================

def reset_day_state():
    st.session_state.bastao_queue = []
    st.session_state.status_texto = {n: 'IndisponÃ­vel' for n in CONSULTORES}
    st.session_state.daily_logs = []
    st.session_state.report_last_run_date = get_brazil_time()

def ensure_daily_reset():
    now_br = get_brazil_time()
    last_run = st.session_state.report_last_run_date
    if isinstance(last_run, str):
        try: last_run_dt = datetime.fromisoformat(last_run).date()
        except: last_run_dt = date.min
    elif isinstance(last_run, datetime):
        last_run_dt = last_run.date()
    else:
        last_run_dt = date.min

    if now_br.date() > last_run_dt:
        if st.session_state.daily_logs: 
            send_daily_report_to_webhook()
            full_state = {
                'date': now_br.isoformat(),
                'logs': st.session_state.daily_logs,
                'queue_final': st.session_state.bastao_queue
            }
            send_state_dump_webhook(full_state)
        reset_day_state()
        save_state()

def auto_manage_time():
    ensure_daily_reset()

# ============================================
# 6. LOGICA DO SISTEMA (DOCUMENTOS E WEBHOOKS)
# ============================================
def verificar_duplicidade_certidao(tipo, processo=None, data=None):
    sb = get_supabase()
    if not sb: return False
    try:
        query = sb.table("certidoes_registro").select("*")
        if tipo in ['FÃ­sico', 'EletrÃ´nico', 'FÃ­sica', 'EletrÃ´nica'] and processo:
            proc_limpo = str(processo).strip()
            if not proc_limpo: return False
            response = query.eq("processo", proc_limpo).execute()
            return len(response.data) > 0
        return False
    except Exception as e:
        return False

def salvar_certidao_db(dados):
    sb = get_supabase()
    if not sb: return False
    try:
        if isinstance(dados.get('data'), (date, datetime)):
            dados['data'] = dados['data'].isoformat()
        if 'hora_periodo' in dados:
             if dados['hora_periodo']:
                dados['motivo'] = f"{dados.get('motivo', '')} - Hora/PerÃ­odo: {dados['hora_periodo']}"
             del dados['hora_periodo']
        if 'n_processo' in dados: 
            dados['processo'] = dados.pop('n_processo')
        if 'n_chamado' in dados: 
            dados['incidente'] = dados.pop('n_chamado')
        if 'data_evento' in dados:
            dados['data'] = dados.pop('data_evento')

        sb.table("certidoes_registro").insert(dados).execute()
        return True
    except Exception as e:
        st.error(f"Erro ao salvar no Supabase: {e}")
        return False

def gerar_docx_certidao_internal(tipo, numero, data, consultor, motivo, chamado="", hora="", nome_parte=""):
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0); section.right_margin = Cm(3.0)
        style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
        
        # CabeÃ§alho
        head_p = doc.add_paragraph(); head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = head_p.add_run("TRIBUNAL DE JUSTIÃ‡A DO ESTADO DE MINAS GERAIS\n")
        runner.bold = True
        head_p.add_run("Rua Ouro Preto, NÂ° 1564 - Bairro Santo Agostinho - CEP 30170-041 - Belo Horizonte - MG\nwww.tjmg.jus.br - Andar: 3Âº e 4Âº PV")
        doc.add_paragraph("\n")
        
        # NumeraÃ§Ã£o
        if tipo == 'Geral': p_num = doc.add_paragraph(f"Parecer GEJUD/DIRTEC/TJMG nÂº ____/2026. Assunto: Notifica erro no â€œJPe â€“ 2Âª InstÃ¢nciaâ€ ao peticionar.")
        else: p_num = doc.add_paragraph(f"Parecer TÃ©cnico GEJUD/DIRTEC/TJMG nÂº ____/2026. Assunto: Notifica erro no â€œJPe â€“ 2Âª InstÃ¢nciaâ€ ao peticionar.")
        p_num.runs[0].bold = True
        
        # Data
        data_extenso_str = ""
        try:
            dt_obj = datetime.strptime(data, "%d/%m/%Y")
            meses = {1:'janeiro', 2:'fevereiro', 3:'marÃ§o', 4:'abril', 5:'maio', 6:'junho', 
                     7:'julho', 8:'agosto', 9:'setembro', 10:'outubro', 11:'novembro', 12:'dezembro'}
            data_extenso_str = f"Belo Horizonte, {dt_obj.day} de {meses[dt_obj.month]} de {dt_obj.year}"
        except:
            data_extenso_str = f"Belo Horizonte, {data}" 
            
        doc.add_paragraph(data_extenso_str)
        doc.add_paragraph(f"Exmo(a). Senhor(a) Relator(a),")
        
        if tipo == 'Geral':
            corpo = doc.add_paragraph(); corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            txt = (f"Para fins de cumprimento dos artigos 13 e 14 da ResoluÃ§Ã£o nÂº 780/2014 do Tribunal de JustiÃ§a do Estado de Minas Gerais, "
                   f"informamos que em {data} houve indisponibilidade do portal JPe, superior a uma hora, {hora}, que impossibilitou o peticionamento eletrÃ´nico de recursos em processos que jÃ¡ tramitavam no sistema.")
            corpo.add_run(txt)
            doc.add_paragraph("\nColocamo-nos Ã  disposiÃ§Ã£o para outras que se fizerem necessÃ¡rias.")
            
        elif tipo in ['EletrÃ´nica', 'EletrÃ´nico']:
            corpo = doc.add_paragraph(); corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo.add_run(f"Informamos que de {data}, houve indisponibilidade especÃ­fica do sistema para o peticionamento do processo nÂº {numero}")
            if nome_parte: corpo.add_run(f", Parte/Advogado: {nome_parte}")
            corpo.add_run(".\n\n")
            corpo.add_run(f"O Chamado de nÃºmero {chamado if chamado else '_____'}, foi aberto e encaminhado Ã  DIRTEC (Diretoria Executiva de Tecnologia da InformaÃ§Ã£o e ComunicaÃ§Ã£o).\n\n")
            corpo.add_run("Esperamos ter prestado as informaÃ§Ãµes solicitadas e colocamo-nos Ã  disposiÃ§Ã£o para outras que se fizerem necessÃ¡rias.")

        elif tipo in ['FÃ­sica', 'FÃ­sico']:
            corpo1 = doc.add_paragraph(); corpo1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo1.add_run(f"Informamos que no dia {data}, houve indisponibilidade especÃ­fica do sistema para o peticionamento do processo nÂº {numero}")
            if nome_parte: corpo1.add_run(f", Parte/Advogado: {nome_parte}")
            corpo1.add_run(".")
            
            corpo2 = doc.add_paragraph(); corpo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo2.add_run(f"O Chamado de nÃºmero {chamado if chamado else '_____'}, foi aberto e encaminhado Ã  DIRTEC (Diretoria Executiva de Tecnologia da InformaÃ§Ã£o e ComunicaÃ§Ã£o).")
            
            corpo3 = doc.add_paragraph(); corpo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo3.add_run("Diante da indisponibilidade especÃ­fica, nÃ£o havendo um prazo para soluÃ§Ã£o do problema, a Primeira Vice-PresidÃªncia recomenda o ingresso dos autos fÃ­sicos, nos termos do Â§ 2Âº, do artigo 14Âº, da ResoluÃ§Ã£o nÂº 780/2014, do Tribunal de JustiÃ§a do Estado de Minas Gerais.")
            doc.add_paragraph("Colocamo-nos Ã  disposiÃ§Ã£o para outras informaÃ§Ãµes que se fizerem necessÃ¡rias.")

        doc.add_paragraph("\nRespeitosamente,")
        sign = doc.add_paragraph("\n___________________________________\nWaner Andrade Silva\n0-009020-9\nCoordenaÃ§Ã£o de AnÃ¡lise e IntegraÃ§Ã£o de Sistemas Judiciais Informatizados - COJIN\nGerÃªncia de Sistemas Judiciais - GEJUD\nDiretoria Executiva de Tecnologia da InformaÃ§Ã£o e ComunicaÃ§Ã£o - DIRTEC")
        sign.runs[0].bold = True 
        
        buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
        return buffer
    except: return None

# Webhooks
def send_chat_notification_internal(consultor, status):
    if CHAT_WEBHOOK_BASTAO and status == 'BastÃ£o':
        msg = f"ğŸ‰ **BASTÃƒO GIRADO!** ğŸ‰ \n\n- **Novo(a) ResponsÃ¡vel:** {consultor}\n- **Acesse o Painel:** {APP_URL_CLOUD}"
        try: send_to_chat("bastao", msg); return True
        except: return False
    return False

def send_state_dump_webhook(state_data):
    if not WEBHOOK_STATE_DUMP: return False
    try:
        sanitized_data = clean_data_for_db(state_data)
        headers = {'Content-Type': 'application/json'}
        requests.post(WEBHOOK_STATE_DUMP, data=json.dumps(sanitized_data), headers=headers, timeout=5)
        return True
    except: return False

def send_horas_extras_to_chat(consultor, data, inicio, tempo, motivo):
    msg = f"â° **Registro de Horas Extras**\n\nğŸ‘¤ **Consultor:** {consultor}\nğŸ“… **Data:** {data.strftime('%d/%m/%Y')}\nğŸ• **InÃ­cio:** {inicio.strftime('%H:%M')}\nâ±ï¸ **Tempo Total:** {tempo}\nğŸ“ **Motivo:** {motivo}"
    try: send_to_chat("extras", msg); return True
    except: return False

def send_atendimento_to_chat(consultor, data, usuario, nome_setor, sistema, descricao, canal, desfecho, jira_opcional=""):
    jira_str = f"\nğŸ”¢ **Jira:** CESUPE-{jira_opcional}" if jira_opcional else ""
    msg = f"ğŸ“‹ **Novo Registro de Atendimento**\n\nğŸ‘¤ **Consultor:** {consultor}\nğŸ“… **Data:** {data.strftime('%d/%m/%Y')}\nğŸ‘¥ **UsuÃ¡rio:** {usuario}\nğŸ¢ **Nome/Setor:** {nome_setor}\nğŸ’» **Sistema:** {sistema}\nğŸ“ **DescriÃ§Ã£o:** {descricao}\nğŸ“ **Canal:** {canal}\nâœ… **Desfecho:** {desfecho}{jira_str}"
    try: send_to_chat("registro", msg); return True
    except: return False

def send_chamado_to_chat(consultor, texto):
    if not consultor or consultor == 'Selecione um nome' or not texto.strip(): return False
    data_envio = get_brazil_time().strftime('%d/%m/%Y %H:%M')
    msg = f"ğŸ†˜ **Rascunho de Chamado/Jira**\nğŸ“… **Data:** {data_envio}\n\nğŸ‘¤ **Autor:** {consultor}\n\nğŸ“ **Texto:**\n{texto}"
    try: send_to_chat('chamado', msg); return True
    except:
        try: send_to_chat('registro', msg); return True
        except: return False

def handle_erro_novidade_submission(consultor, titulo, objetivo, relato, resultado):
    data_envio = get_brazil_time().strftime("%d/%m/%Y %H:%M")
    msg = f"ğŸ› **Novo Relato de Erro/Novidade**\nğŸ“… **Data:** {data_envio}\n\nğŸ‘¤ **Autor:** {consultor}\nğŸ“Œ **TÃ­tulo:** {titulo}\n\nğŸ¯ **Objetivo:**\n{objetivo}\n\nğŸ§ª **Relato:**\n{relato}\n\nğŸ **Resultado:**\n{resultado}"
    try: send_to_chat("erro", msg); return True
    except: return False

def send_sessao_to_chat_fn(consultor, texto_mensagem):
    return True

def handle_sugestao_submission(consultor, texto):
    data_envio = get_brazil_time().strftime("%d/%m/%Y %H:%M")
    ip_usuario = get_remote_ip()
    msg = f"ğŸ’¡ **Nova SugestÃ£o**\nğŸ“… **Data:** {data_envio}\nğŸ‘¤ **Autor:** {consultor}\nğŸŒ **IP:** {ip_usuario}\n\nğŸ“ **SugestÃ£o:**\n{texto}"
    try: send_to_chat("extras", msg); return True
    except: return False

# ============================================
# 7. FUNÃ‡Ã•ES DE ESTADO E LÃ“GICA (REORGANIZADAS)
# ============================================

def save_state():
    try:
        last_run = st.session_state.report_last_run_date
        visual_queue_calculated = get_ordered_visual_queue(st.session_state.bastao_queue, st.session_state.status_texto)
        
        state_to_save = {
            'status_texto': st.session_state.status_texto, 'bastao_queue': st.session_state.bastao_queue,
            'visual_queue': visual_queue_calculated, 'skip_flags': st.session_state.skip_flags, 
            'current_status_starts': st.session_state.current_status_starts,
            'bastao_counts': st.session_state.bastao_counts, 'priority_return_queue': st.session_state.priority_return_queue,
            'bastao_start_time': st.session_state.bastao_start_time, 
            'report_last_run_date': last_run, 
            'rotation_gif_start_time': st.session_state.get('rotation_gif_start_time'), 
            'auxilio_ativo': st.session_state.get('auxilio_ativo', False), 
            'daily_logs': st.session_state.daily_logs, 
            'simon_ranking': st.session_state.get('simon_ranking', []),
            'previous_states': st.session_state.get('previous_states', {})
        }
        # Limpeza Ã© feita dentro de save_state_to_db agora, mas podemos chamar aqui tambÃ©m para garantir
        save_state_to_db(state_to_save)
        
        # INVALIDAÃ‡ÃƒO DE CACHE (CRÃTICO)
        load_state_from_db.clear()
        
    except Exception as e: print(f"Erro save: {e}")

def format_time_duration(duration):
    if not isinstance(duration, timedelta): return '--:--:--'
    s = int(duration.total_seconds()); h, s = divmod(s, 3600); m, s = divmod(s, 60)
    return f'{h:02}:{m:02}:{s:02}'

def sync_state_from_db():
    try:
        # Usa a funÃ§Ã£o com cache
        db_data = load_state_from_db()
        if not db_data: return
        keys = ['status_texto', 'bastao_queue', 'skip_flags', 'bastao_counts', 'priority_return_queue', 'daily_logs', 'simon_ranking', 'previous_states']
        for k in keys:
            if k in db_data: 
                # PaginaÃ§Ã£o
                if k == 'daily_logs' and isinstance(db_data[k], list) and len(db_data[k]) > 150:
                    st.session_state[k] = db_data[k][-150:] 
                else:
                    st.session_state[k] = db_data[k]
                    
        if 'bastao_start_time' in db_data and db_data['bastao_start_time']:
            try:
                if isinstance(db_data['bastao_start_time'], str): st.session_state['bastao_start_time'] = datetime.fromisoformat(db_data['bastao_start_time'])
                else: st.session_state['bastao_start_time'] = db_data['bastao_start_time']
            except: pass
        if 'current_status_starts' in db_data:
            starts = db_data['current_status_starts']
            for nome, val in starts.items():
                if isinstance(val, str):
                    try: st.session_state.current_status_starts[nome] = datetime.fromisoformat(val)
                    except: pass
                else: st.session_state.current_status_starts[nome] = val
    except Exception as e: print(f"Erro sync: {e}")

def log_status_change(consultor, old_status, new_status, duration):
    if not isinstance(duration, timedelta): duration = timedelta(0)
    now_br = get_brazil_time()
    st.session_state.daily_logs.append({
        'timestamp': now_br, 'consultor': consultor, 
        'old_status': old_lbl if 'old_lbl' in locals() else old_status or 'Fila', 
        'new_status': new_lbl if 'new_lbl' in locals() else new_status or 'Fila', 
        'duration': duration, 'ip': st.session_state.get('device_id_val', 'unknown')
    })
    
    # Limite local de 150 para nÃ£o estourar RAM
    if len(st.session_state.daily_logs) > 150:
        st.session_state.daily_logs = st.session_state.daily_logs[-150:]
        
    st.session_state.current_status_starts[consultor] = now_br

# --- LOGICA DE STATUS (FUNÃ‡Ã•ES HELPERS) ---
def find_next_holder_index(current_index, queue, skips):
    if not queue: return -1
    n = len(queue); start_index = (current_index + 1) % n
    for i in range(n):
        idx = (start_index + i) % n
        if not skips.get(queue[idx], False): return idx
    if n > 1:
        proximo_imediato_idx = (current_index + 1) % n
        nome_escolhido = queue[proximo_imediato_idx]; st.session_state.skip_flags[nome_escolhido] = False 
        return proximo_imediato_idx
    return -1

def send_daily_report_to_webhook():
    logs = st.session_state.daily_logs
    if not logs: return False
    # Envio simplificado para evitar erro
    try: send_state_dump_webhook({'logs': st.session_state.daily_logs})
    except: pass

def check_and_assume_baton(forced_successor=None, immune_consultant=None):
    queue, skips = st.session_state.bastao_queue, st.session_state.skip_flags
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
    is_valid = (current_holder and current_holder in queue)
    target = forced_successor if forced_successor else (current_holder if is_valid else None)
    
    if not target:
        curr_idx = queue.index(current_holder) if (current_holder and current_holder in queue) else -1
        idx = find_next_holder_index(curr_idx, queue, skips)
        target = queue[idx] if idx != -1 else None
        
    changed = False; now = get_brazil_time()
    
    for c in CONSULTORES:
        if c != immune_consultant: 
            if c != target and 'BastÃ£o' in st.session_state.status_texto.get(c, ''):
                log_status_change(c, 'BastÃ£o', 'IndisponÃ­vel', now - st.session_state.current_status_starts.get(c, now))
                st.session_state.status_texto[c] = 'IndisponÃ­vel'; changed = True
    
    if target:
        curr_s = st.session_state.status_texto.get(target, '')
        if 'BastÃ£o' not in curr_s:
            old_s = curr_s; new_s = f"BastÃ£o | {old_s}" if old_s and old_s != "IndisponÃ­vel" else "BastÃ£o"
            log_status_change(target, old_s, new_s, now - st.session_state.current_status_starts.get(target, now))
            st.session_state.status_texto[target] = new_s; st.session_state.bastao_start_time = now
            if current_holder != target: st.session_state.play_sound = True; send_chat_notification_internal(target, 'BastÃ£o')
            st.session_state.skip_flags[target] = False
            changed = True
    elif not target and current_holder:
        if current_holder != immune_consultant:
            log_status_change(current_holder, 'BastÃ£o', 'IndisponÃ­vel', now - st.session_state.current_status_starts.get(current_holder, now))
            st.session_state.status_texto[current_holder] = 'IndisponÃ­vel'; changed = True
            
    if changed: save_state()
    return changed

# --- AÃ‡Ã•ES PRINCIPAIS (QUE CHAMAM AS HELPERS) ---
def update_status(novo_status: str, marcar_indisponivel: bool = False, manter_fila_atual: bool = False):
    selected = st.session_state.get('consultor_selectbox')
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    
    ensure_daily_reset()
    now_br = get_brazil_time()
    current = st.session_state.status_texto.get(selected, '')
    forced_successor = None
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in (s or '')), None)
    
    if novo_status == 'AlmoÃ§o':
        st.session_state.previous_states[selected] = {
            'status': current,
            'in_queue': selected in st.session_state.bastao_queue
        }
    
    if marcar_indisponivel:
        st.session_state[f'check_{selected}'] = False; st.session_state.skip_flags[selected] = True
        if selected in st.session_state.bastao_queue:
            if selected == current_holder:
                idx = st.session_state.bastao_queue.index(selected)
                nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
                if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
            st.session_state.bastao_queue.remove(selected)
            if selected not in st.session_state.priority_return_queue: st.session_state.priority_return_queue.append(selected)
    
    if novo_status == 'IndisponÃ­vel':
        if selected in st.session_state.bastao_queue:
            st.session_state.bastao_queue.remove(selected)

    elif not manter_fila_atual:
        if selected not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(selected)
        st.session_state[f'check_{selected}'] = True
        st.session_state.skip_flags[selected] = False
    
    final_status = (novo_status or '').strip()
    if selected == current_holder and selected in st.session_state.bastao_queue:
          final_status = ('BastÃ£o | ' + final_status).strip(' |') if final_status else 'BastÃ£o'
    if not final_status and (selected not in st.session_state.bastao_queue): final_status = 'IndisponÃ­vel'
    
    log_status_change(selected, current, final_status, now - st.session_state.current_status_starts.get(selected, now))
    st.session_state.status_texto[selected] = final_status
    check_and_assume_baton(forced_successor)
    save_state()

def toggle_queue(consultor):
    ensure_daily_reset(); st.session_state.gif_warning = False; now_br = get_brazil_time()
    if consultor in st.session_state.bastao_queue:
        current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
        forced_successor = None
        if consultor == current_holder:
            idx = st.session_state.bastao_queue.index(consultor)
            nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
            if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
        st.session_state.bastao_queue.remove(consultor)
        st.session_state.status_texto[consultor] = 'IndisponÃ­vel'
        check_and_assume_baton(forced_successor)
    else:
        st.session_state.bastao_queue.append(consultor)
        st.session_state.skip_flags[consultor] = False
        if consultor in st.session_state.priority_return_queue: st.session_state.priority_return_queue.remove(consultor)
        st.session_state.status_texto[consultor] = ''
        check_and_assume_baton()
    save_state()

def rotate_bastao():
    ensure_daily_reset(); selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    
    queue = st.session_state.bastao_queue; skips = st.session_state.skip_flags
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
    
    # === VALIDAÃ‡ÃƒO: SÃ“ TITULAR PODE PASSAR ===
    if selected != current_holder:
        st.error(f"âš ï¸ Apenas quem estÃ¡ com o bastÃ£o ({current_holder}) pode passÃ¡-lo!")
        return
        
    current_index = queue.index(current_holder) if current_holder in queue else -1
    if current_index == -1: check_and_assume_baton(); return
    next_idx = find_next_holder_index(current_index, queue, skips)
    if next_idx == -1 and len(queue) > 1: next_idx = (current_index + 1) % len(queue)
    if next_idx != -1:
        n_queue = len(queue); tmp_idx = (current_index + 1) % n_queue
        while tmp_idx != next_idx:
            skipped_name = queue[tmp_idx]
            if st.session_state.skip_flags.get(skipped_name, False): st.session_state.skip_flags[skipped_name] = False
            tmp_idx = (tmp_idx + 1) % n_queue
        next_holder = queue[next_idx]; st.session_state.skip_flags[next_holder] = False; now_br = get_brazil_time()
        old_h_status = st.session_state.status_texto[current_holder]
        new_h_status = old_h_status.replace('BastÃ£o | ', '').replace('BastÃ£o', '').strip()
        log_status_change(current_holder, old_h_status, new_h_status, now_br - (st.session_state.bastao_start_time or now_br))
        st.session_state.status_texto[current_holder] = new_h_status
        old_n_status = st.session_state.status_texto.get(next_holder, '')
        new_n_status = f"BastÃ£o | {old_n_status}" if old_n_status else "BastÃ£o"
        log_status_change(next_holder, old_n_status, new_n_status, timedelta(0))
        st.session_state.status_texto[next_holder] = new_n_status
        st.session_state.bastao_start_time = now_br
        st.session_state.bastao_counts[current_holder] = st.session_state.bastao_counts.get(current_holder, 0) + 1
        st.session_state.play_sound = True; send_chat_notification_internal(next_holder, 'BastÃ£o')
        save_state()
    else: st.warning('NinguÃ©m elegÃ­vel.'); check_and_assume_baton()

def toggle_skip():
    selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    if selected not in st.session_state.bastao_queue: st.warning(f'{selected} nÃ£o estÃ¡ na fila do bastÃ£o.'); return
    novo = not st.session_state.skip_flags.get(selected, False)
    st.session_state.skip_flags[selected] = novo
    save_state()

def toggle_presence_btn():
    selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    toggle_queue(selected)

def enter_from_indisponivel(c):
    if c not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(c)
    st.session_state.status_texto[c] = ''; save_state()

def toggle_view(v):
    if st.session_state.active_view == v: st.session_state.active_view = None
    else: st.session_state.active_view = v

def reset_day_state():
    st.session_state.bastao_queue = []; st.session_state.status_texto = {n: 'IndisponÃ­vel' for n in CONSULTORES}
    st.session_state.daily_logs = []; st.session_state.report_last_run_date = get_brazil_time()

def ensure_daily_reset():
    now_br = get_brazil_time(); last_run = st.session_state.report_last_run_date
    if isinstance(last_run, str):
        try: last_run_dt = datetime.fromisoformat(last_run).date()
        except: last_run_dt = date.min
    elif isinstance(last_run, datetime):
        last_run_dt = last_run.date()
    else:
        last_run_dt = date.min

    if now_br.date() > last_run_dt:
        if st.session_state.daily_logs: 
            send_daily_report_to_webhook()
            full_state = {
                'date': now_br.isoformat(),
                'logs': st.session_state.daily_logs,
                'queue_final': st.session_state.bastao_queue
            }
            send_state_dump_webhook(full_state)
        reset_day_state()
        save_state()

def auto_manage_time():
    ensure_daily_reset()

def init_session_state():
    dev = get_browser_id(); 
    if dev: st.session_state['device_id_val'] = dev
    if 'db_loaded' not in st.session_state:
        db = load_state_from_db()
        if 'report_last_run_date' in db and isinstance(db['report_last_run_date'], str):
            try: db['report_last_run_date'] = datetime.fromisoformat(db['report_last_run_date'])
            except: db['report_last_run_date'] = datetime.min
        st.session_state.update(db); st.session_state['db_loaded'] = True
    
    defaults = {
        'bastao_start_time': None, 'report_last_run_date': datetime.min, 'rotation_gif_start_time': None,
        'play_sound': False, 'gif_warning': False, 'lunch_warning_info': None, 'last_reg_status': None,
        'chamado_guide_step': 0, 'auxilio_ativo': False, 'active_view': None,
        'consultor_selectbox': "Selecione um nome", 'status_texto': {n: 'IndisponÃ­vel' for n in CONSULTORES},
        'bastao_queue': [], 'skip_flags': {}, 'current_status_starts': {n: get_brazil_time() for n in CONSULTORES},
        'bastao_counts': {n: 0 for n in CONSULTORES}, 'priority_return_queue': [], 'daily_logs': [], 'simon_ranking': [],
        'word_buffer': None, 'aviso_duplicidade': False, 'previous_states': {}, 'view_logmein_ui': False,
        'last_cleanup': time.time(), 'last_hard_cleanup': time.time()
    }
    for k, v in defaults.items():
        if k not in st.session_state: st.session_state[k] = v
    for n in CONSULTORES:
        st.session_state.skip_flags.setdefault(n, False); st.session_state[f'check_{n}'] = n in st.session_state.bastao_queue

def open_logmein_ui(): st.session_state.view_logmein_ui = True
def close_logmein_ui(): st.session_state.view_logmein_ui = False

# ============================================
# 8. INTERFACE
# ============================================
st.set_page_config(page_title="Controle Cesupe 2026", layout="wide", page_icon="ğŸ­")
st.markdown("""<style>div.stButton > button {width: 100%; height: 3rem;}</style>""", unsafe_allow_html=True)

init_session_state(); memory_sweeper(); auto_manage_time()

st.components.v1.html("<script>window.scrollTo(0, 0);</script>", height=0)
st.info("ğŸ—ï¸ Fevereiro Laranja Ã© um convite Ã  consciÃªncia e Ã  aÃ§Ã£o: ele chama atenÃ§Ã£o para a leucemia e para a importÃ¢ncia do diagnÃ³stico precoce, que pode salvar vidas. ğŸ’›ğŸ§¡ Informar, apoiar quem estÃ¡ em tratamento e incentivar a doaÃ§Ã£o de sangue e de medula Ã³ssea sÃ£o atitudes que fazem diferenÃ§a. Compartilhe, converse e, se puder, cadastre-se como doador â€” um gesto simples pode ser a esperanÃ§a de alguÃ©m.")

# ----------------------------------------------------
# FRAGMENTO DE VISUALIZAÃ‡ÃƒO: HEADER + FILA
# ----------------------------------------------------
@st.fragment(run_every=15)
def render_header_info_left():
    sync_state_from_db()
    
    # 1. Header
    c_topo_esq, c_topo_dir = st.columns([2, 1], vertical_alignment="bottom")
    with c_topo_esq:
        img = get_img_as_base64_cached(PUG2026_FILENAME); src = f"data:image/png;base64,{img}" if img else GIF_BASTAO_HOLDER
        st.markdown(f"""<div style="display: flex; align-items: center; gap: 15px;"><h1 style="margin: 0; padding: 0; font-size: 2.2rem; color: #FF8C00; text-shadow: 1px 1px 2px #FF4500;">Controle BastÃ£o Cesupe 2026 {BASTAO_EMOJI}</h1><img src="{src}" style="width: 150px; height: 150px; border-radius: 10px; border: 4px solid #FF8C00; object-fit: cover;"></div>""", unsafe_allow_html=True)
    with c_topo_dir:
        c_sub1, c_sub2 = st.columns([2, 1], vertical_alignment="bottom")
        with c_sub1: 
             novo_responsavel = st.selectbox("Assumir BastÃ£o (RÃ¡pido)", options=["Selecione"] + CONSULTORES, label_visibility="collapsed", key="quick_enter")
        with c_sub2:
            if st.button("ğŸš€ Entrar", use_container_width=True, key="btn_entrar_header"):
                if novo_responsavel != "Selecione": toggle_queue(novo_responsavel); st.rerun()
        st.caption(f"ID: ...{st.session_state.get('device_id_val', '???')[-4:]}")
        # VÃ¡lvula de Escape
        if st.button("ğŸ”„ Atualizar Agora", use_container_width=True): 
             load_state_from_db.clear(); st.rerun()

    st.markdown("<hr style='border: 1px solid #FF8C00; margin-top: 5px; margin-bottom: 20px;'>", unsafe_allow_html=True)
    
    # 2. Dados da Fila
    queue = st.session_state.bastao_queue
    skips = st.session_state.skip_flags
    responsavel = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
    curr_idx = queue.index(responsavel) if responsavel in queue else -1
    prox_idx = find_next_holder_index(curr_idx, queue, skips)
    proximo = queue[prox_idx] if prox_idx != -1 else None

    # Card ResponsÃ¡vel
    st.header("ResponsÃ¡vel pelo BastÃ£o")
    if responsavel:
        st.markdown(f"""<div style="background: linear-gradient(135deg, #FFF3E0 0%, #FFFFFF 100%); border: 3px solid #FF8C00; padding: 25px; border-radius: 15px; display: flex; align-items: center; box-shadow: 0 4px 15px rgba(255, 140, 0, 0.3); margin-bottom: 20px;"><div style="flex-shrink: 0; margin-right: 25px;"><img src="{GIF_BASTAO_HOLDER}" style="width: 90px; height: 90px; border-radius: 50%; object-fit: cover; border: 2px solid #FF8C00;"></div><div><span style="font-size: 14px; color: #555; font-weight: bold; text-transform: uppercase; letter-spacing: 1.5px;">Atualmente com:</span><br><span style="font-size: 42px; font-weight: 800; color: #FF4500; line-height: 1.1;">{responsavel}</span></div></div>""", unsafe_allow_html=True)
        dur = get_brazil_time() - (st.session_state.bastao_start_time or get_brazil_time())
        st.caption(f"â±ï¸ Tempo com o bastÃ£o: **{format_time_duration(dur)}**")
    else: st.markdown('<h2>(NinguÃ©m com o bastÃ£o)</h2>', unsafe_allow_html=True)
    
    # Texto da Fila
    st.markdown("###"); st.header("PrÃ³ximos da Fila")
    if responsavel and responsavel in queue:
        c_idx = queue.index(responsavel)
        raw_ordered = queue[c_idx+1:] + queue[:c_idx]
    else: raw_ordered = list(queue)
    lista_pularam = [n for n in queue if skips.get(n, False) and n != responsavel]
    demais_na_fila = [n for n in raw_ordered if n != proximo and not skips.get(n, False)]
    
    if proximo: st.markdown(f"**PrÃ³ximo BastÃ£o:** {proximo}")
    else: st.markdown("**PrÃ³ximo BastÃ£o:** _NinguÃ©m elegÃ­vel_")
    if demais_na_fila: st.markdown(f"**Demais na fila:** {', '.join(demais_na_fila)}")
    else: st.markdown("**Demais na fila:** _Vazio_")
    if lista_pularam: st.markdown(f"**Consultor(es) pulou(pularam) o bastÃ£o:** {', '.join(lista_pularam)}")

# FRAGMENTO DIREITO: LISTA DE STATUS
@st.fragment(run_every=15)
def render_status_list():
    sync_state_from_db()
    queue = st.session_state.bastao_queue
    skips = st.session_state.skip_flags
    responsavel = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)

    st.header('Status dos(as) Consultores(as)')
    ui_lists = {'fila': [], 'almoco': [], 'saida': [], 'ausente': [], 'atividade_especifica': [], 'sessao_especifica': [], 'projeto_especifico': [], 'reuniao_especifica': [], 'treinamento_especifico': [], 'indisponivel': [], 'presencial_especifico': []}
    for nome in CONSULTORES:
        if nome in st.session_state.bastao_queue: ui_lists['fila'].append(nome)
        status = st.session_state.status_texto.get(nome, 'IndisponÃ­vel'); status = status if status is not None else 'IndisponÃ­vel'
        if status in ('', None): pass
        elif status == 'AlmoÃ§o': ui_lists['almoco'].append(nome)
        elif status == 'SaÃ­da rÃ¡pida': ui_lists['saida'].append(nome)
        elif status == 'IndisponÃ­vel' and nome not in st.session_state.bastao_queue: ui_lists['indisponivel'].append(nome)
        if isinstance(status, str):
            if 'SessÃ£o:' in status or status.strip() == 'SessÃ£o': ui_lists['sessao_especifica'].append((nome, status.replace('SessÃ£o:', '').strip()))
            if 'ReuniÃ£o:' in status or status.strip() == 'ReuniÃ£o': ui_lists['reuniao_especifica'].append((nome, status.replace('ReuniÃ£o:', '').strip()))
            if 'Projeto:' in status or status.strip() == 'Projeto': ui_lists['projeto_especifico'].append((nome, status.replace('Projeto:', '').strip()))
            if 'Treinamento:' in status or status.strip() == 'Treinamento': ui_lists['treinamento_especifico'].append((nome, status.replace('Treinamento:', '').strip()))
            if 'Atividade:' in status or status.strip() == 'Atendimento': ui_lists['atividade_especifica'].append((nome, status.replace('Atividade:', '').strip()))
            if 'Atendimento Presencial:' in status: ui_lists['presencial_especifico'].append((nome, status.replace('Atendimento Presencial:', '').strip()))

    st.subheader(f'âœ… Na Fila ({len(ui_lists["fila"])})')
    render_order = get_ordered_visual_queue(queue, st.session_state.status_texto)
    if not render_order and queue: render_order = list(queue)
    if not render_order: st.markdown('_NinguÃ©m na fila._')
    else:
        for i, nome in enumerate(render_order):
            if nome not in ui_lists['fila']: continue
            col_nome, col_check = st.columns([0.85, 0.15], vertical_alignment='center')
            col_check.checkbox(' ', key=f'chk_fila_{nome}_frag', value=True, disabled=True, label_visibility='collapsed')
            skip_flag = skips.get(nome, False); status_atual = st.session_state.status_texto.get(nome, '') or ''; extra = ''
            if 'Atividade' in status_atual: extra += ' ğŸ“‹'
            if 'Projeto' in status_atual: extra += ' ğŸ—ï¸'
            if nome == responsavel: display = f'<span style="background-color: #FF8C00; color: #FFF; padding: 2px 6px; border-radius: 5px; font-weight: 800;">ğŸ­ {nome}</span>'
            elif skip_flag: display = f'<strong>{i}Âº {nome}</strong>{extra} <span style="background-color: #FFECB3; padding: 2px 8px; border-radius: 10px;">Pulando â­ï¸</span>'
            else: display = f'<strong>{i}Âº {nome}</strong>{extra} <span style="background-color: #FFE0B2; padding: 2px 8px; border-radius: 10px;">Aguardando</span>'
            col_nome.markdown(display, unsafe_allow_html=True)
    st.markdown('---')

    def _render_section(titulo, icon, itens, cor, key_rm):
        colors = {'orange': '#FFECB3', 'blue': '#BBDEFB', 'teal': '#B2DFDB', 'violet': '#E1BEE7', 'green': '#C8E6C9', 'red': '#FFCDD2', 'grey': '#EEEEEE', 'yellow': '#FFF9C4'}
        bg_hex = colors.get(cor, '#EEEEEE'); st.subheader(f'{icon} {titulo} ({len(itens)})')
        if not itens: st.markdown(f'_Nenhum._')
        else:
            for item in itens:
                nome = item[0] if isinstance(item, tuple) else item
                desc = item[1] if isinstance(item, tuple) else titulo
                col_n, col_c = st.columns([0.85, 0.15], vertical_alignment='center')
                if titulo == 'IndisponÃ­vel': 
                    if col_c.checkbox(' ', key=f'chk_{titulo}_{nome}_frag', value=False, label_visibility='collapsed'):
                        enter_from_indisponivel(nome); st.rerun()
                col_n.markdown(f"<div style='font-size: 16px; margin: 2px 0;'><strong>{nome}</strong><span style='background-color: {bg_hex}; color: #333; padding: 2px 8px; border-radius: 12px; font-size: 14px; margin-left: 8px;'>{desc}</span></div>", unsafe_allow_html=True)
        st.markdown('---')
        
    _render_section('Atend. Presencial', 'ğŸ¤', ui_lists['presencial_especifico'], 'yellow', 'Atendimento Presencial')
    _render_section('Em Demanda', 'ğŸ“‹', ui_lists['atividade_especifica'], 'orange', 'Atividade')
    _render_section('Projetos', 'ğŸ—ï¸', ui_lists['projeto_especifico'], 'blue', 'Projeto')
    _render_section('Treinamento', 'ğŸ“', ui_lists['treinamento_especifico'], 'teal', 'Treinamento')
    _render_section('ReuniÃµes', 'ğŸ“…', ui_lists['reuniao_especifica'], 'violet', 'ReuniÃ£o')
    _render_section('AlmoÃ§o', 'ğŸ½ï¸', ui_lists['almoco'], 'red', 'AlmoÃ§o')
    _render_section('SessÃ£o', 'ğŸ™ï¸', ui_lists['sessao_especifica'], 'green', 'SessÃ£o')
    _render_section('SaÃ­da rÃ¡pida', 'ğŸš¶', ui_lists['saida'], 'red', 'SaÃ­da rÃ¡pida')
    _render_section('IndisponÃ­vel', 'âŒ', ui_lists['indisponivel'], 'grey', '')

# =========================================================================
# LAYOUT PRINCIPAL (ONDE A MÃGICA ACONTECE)
# =========================================================================

# Criamos as colunas FORA dos fragmentos para poder injetar conteÃºdo nelas
col_principal, col_disponibilidade = st.columns([1.5, 1])

# 1. Renderiza o conteÃºdo que se atualiza sozinho
with col_principal:
    # Apenas se nÃ£o houver menu ativo, renderiza o topo automÃ¡tico
    render_header_info_left()

with col_disponibilidade:
    render_status_list()

# 2. Renderiza os BotÃµes de AÃ§Ã£o na Coluna Esquerda (FORA do fragmento para funcionar sempre)
with col_principal:
    st.markdown("### ğŸ® Painel de AÃ§Ã£o")
    c_nome, c_act1, c_act2, c_act3 = st.columns([2, 1, 1, 1], vertical_alignment="bottom")
    with c_nome:
        st.selectbox('Selecione:', ['Selecione um nome'] + CONSULTORES, key='consultor_selectbox', label_visibility='collapsed')
    with c_act1:
        st.button("ğŸ­ Entrar/Sair Fila", on_click=toggle_presence_btn, use_container_width=True)
    with c_act2:
        st.button('ğŸ¯ Passar', on_click=rotate_bastao, use_container_width=True)
    with c_act3:
        st.button('â­ï¸ Pular', on_click=toggle_skip, use_container_width=True)
    
    r2c1, r2c2, r2c3, r2c4, r2c5 = st.columns(5)
    r2c1.button('ğŸ“‹ Atividades', on_click=toggle_view, args=('menu_atividades',), use_container_width=True)
    r2c2.button('ğŸ—ï¸ Projeto', on_click=toggle_view, args=('menu_projetos',), use_container_width=True)
    r2c3.button('ğŸ“ Treinamento', on_click=toggle_view, args=('menu_treinamento',), use_container_width=True)
    r2c4.button('ğŸ“… ReuniÃ£o', on_click=toggle_view, args=('menu_reuniao',), use_container_width=True)
    r2c5.button('ğŸ½ï¸ AlmoÃ§o', on_click=update_status, args=('AlmoÃ§o', True), use_container_width=True)
    
    r3c1, r3c2, r3c3, r3c4 = st.columns(4)
    r3c1.button('ğŸ™ï¸ SessÃ£o', on_click=toggle_view, args=('menu_sessao',), use_container_width=True)
    r3c2.button('ğŸš¶ SaÃ­da', on_click=update_status, args=('SaÃ­da rÃ¡pida', True), use_container_width=True)
    r3c3.button('ğŸƒ Sair', on_click=update_status, args=('IndisponÃ­vel', True), use_container_width=True)
    if r3c4.button("ğŸ¤ Atend. Presencial", use_container_width=True): toggle_view('menu_presencial')
    
    st.markdown("####")
    if st.button('ğŸ”‘ LogMeIn', use_container_width=True):
        open_logmein_ui()

    # ============================================
    # LÃ“GICA DE MENUS (APARECEM NA COLUNA PRINCIPAL - INLINE)
    # ============================================

    # LogMeIn UI
    if st.session_state.view_logmein_ui:
        with st.container(border=True):
            st.markdown("### ğŸ’» Acesso LogMeIn")
            l_user, l_in_use = get_logmein_status()
            
            st.image(GIF_LOGMEIN_TARGET, width=180)
            
            if l_in_use:
                st.error(f"ğŸ”´ EM USO POR: **{l_user}**")
                meu_nome = st.session_state.get('consultor_selectbox')
                # Libera se for o dono ou admin
                if meu_nome == l_user or meu_nome in CONSULTORES:
                    if st.button("ğŸ”“ LIBERAR AGORA", type="primary", use_container_width=True):
                        set_logmein_status(None, False)
                        close_logmein_ui()
                        st.rerun()
                    else:
                        st.info("Aguarde a liberaÃ§Ã£o.")
            else:
                st.success("âœ… LIVRE PARA USO")
                meu_nome = st.session_state.get('consultor_selectbox')
                if meu_nome and meu_nome != "Selecione um nome":
                    if st.button("ğŸš€ ASSUMIR AGORA", use_container_width=True):
                        set_logmein_status(meu_nome, True)
                        close_logmein_ui()
                        st.rerun()
                else:
                    st.warning("Selecione seu nome no topo para assumir.")
            
            if st.button("Fechar", use_container_width=True):
                close_logmein_ui()
                st.rerun()

    # --- MENUS DE AÃ‡ÃƒO ---
    if st.session_state.active_view == 'menu_atividades':
        with st.container(border=True):
            at_t = st.multiselect("Tipo:", OPCOES_ATIVIDADES_STATUS); at_e = st.text_input("Detalhe:")
            c1, c2, c3 = st.columns([1, 1, 1])
            with c1:
                if st.button("Confirmar", type="primary", use_container_width=True): 
                    st.session_state.active_view = None
                    update_status(f"Atividade: {', '.join(at_t)} - {at_e}", manter_fila_atual=True) # MANTÃ‰M FILA
            with c2:
                if st.button("Sair de atividades", use_container_width=True):
                    st.session_state.active_view = None
                    update_status("", manter_fila_atual=True) 
            with c3:
                if st.button("Cancelar", use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == 'menu_presencial':
        with st.container(border=True):
            st.subheader('ğŸ¤ Registrar Atendimento Presencial'); local_presencial = st.text_input('Local:', key='pres_local'); objetivo_presencial = st.text_input('Objetivo:', key='pres_obj')
            c_ok, c_cancel = st.columns(2)
            with c_ok:
                if st.button('âœ… Confirmar', type='primary', use_container_width=True):
                    if not local_presencial.strip() or not objetivo_presencial.strip(): st.warning('Preencha Local e Objetivo.')
                    else: st.session_state.active_view = None; update_status(f"Atendimento Presencial: {local_presencial.strip()} - {objetivo_presencial.strip()}", True) # REMOVE DA FILA (TRUE)
            with c_cancel:
                if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == 'menu_projetos':
        with st.container(border=True):
            st.subheader('ğŸ—ï¸ Registrar Projeto')
            proj_nome = st.text_input('Nome do Projeto:', placeholder='Digite o nome do projeto...')
            manter_bastao = st.checkbox("Continuar recebendo bastÃ£o? (Modo Atividade)")
            
            c_ok, c_cancel = st.columns(2)
            with c_ok:
                if st.button('âœ… Confirmar', type='primary', use_container_width=True):
                    if not proj_nome.strip(): st.warning('Digite o nome do projeto.')
                    else: 
                        st.session_state.active_view = None
                        status_msg = f"Projeto: {proj_nome.strip()}"
                        if manter_bastao: update_status(status_msg, manter_fila_atual=True)
                        else: update_status(status_msg, marcar_indisponivel=True)
            with c_cancel:
                if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == 'menu_treinamento':
        with st.container(border=True):
            st.subheader('ğŸ“ Registrar Treinamento'); tema = st.text_input('Tema/ConteÃºdo:'); obs = st.text_input('ObservaÃ§Ã£o (opcional):')
            c_ok, c_cancel = st.columns(2)
            with c_ok:
                if st.button('âœ… Confirmar', type='primary', use_container_width=True):
                    if not tema.strip(): st.warning('Informe o tema.')
                    else: st.session_state.active_view = None; update_status(f"Treinamento: {tema.strip()}" + (f" - {obs.strip()}" if obs.strip() else ""), True)
            with c_cancel:
                if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == 'menu_reuniao':
        with st.container(border=True):
            st.subheader('ğŸ“… Registrar ReuniÃ£o'); assunto = st.text_input('Assunto:'); obs = st.text_input('ObservaÃ§Ã£o (opcional):')
            c_ok, c_cancel = st.columns(2)
            with c_ok:
                if st.button('âœ… Confirmar', type='primary', use_container_width=True):
                    if not assunto.strip(): st.warning('Informe o assunto.')
                    else: st.session_state.active_view = None; update_status(f"ReuniÃ£o: {assunto.strip()}" + (f" - {obs.strip()}" if obs.strip() else ""), True)
            with c_cancel:
                if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == 'menu_sessao':
        with st.container(border=True):
            st.subheader('ğŸ™ï¸ Registrar SessÃ£o')
            sessao_livre = st.text_input('Qual SessÃ£o / CÃ¢mara?'); obs = st.text_input('ObservaÃ§Ã£o (opcional):')
            enviar_chat = st.checkbox('Enviar aviso no chat', value=True)
            c_ok, c_cancel = st.columns(2)
            with c_ok:
                if st.button('âœ… Confirmar', type='primary', use_container_width=True):
                    consultor = st.session_state.get('consultor_selectbox')
                    if not consultor or consultor == 'Selecione um nome': st.error('Selecione um consultor.')
                    elif not sessao_livre.strip(): st.warning('Digite qual a sessÃ£o.')
                    else:
                        st.session_state.active_view = None; update_status(f"SessÃ£o: {sessao_livre}" + (f" - {obs.strip()}" if obs.strip() else ""), True)
            with c_cancel:
                if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()
                    
    # --- FERRAMENTAS ESPECIAIS (ABAIXO DOS BOTÃ•ES, DENTRO DA COLUNA ESQUERDA) ---
    st.markdown("<hr style='border: 1px solid #FF8C00;'>", unsafe_allow_html=True)
    st.markdown("#### Ferramentas")
    
    # DIVIDIDO EM 2 LINHAS PARA CABER OS NOMES
    c_t1, c_t2, c_t3, c_t4 = st.columns(4)
    c_t1.button("ğŸ“‘ Checklist", use_container_width=True, on_click=toggle_view, args=("checklist",))
    c_t2.button("ğŸ†˜ Chamados", use_container_width=True, on_click=toggle_view, args=("chamados",))
    c_t3.button("ğŸ“ Atendimentos", use_container_width=True, on_click=toggle_view, args=("atendimentos",))
    c_t4.button("â° H. Extras", use_container_width=True, on_click=toggle_view, args=("hextras",))
    
    c_t5, c_t6, c_t7 = st.columns(3)
    c_t5.button("ğŸ› Erro/Novidade", use_container_width=True, on_click=toggle_view, args=("erro_novidade",))
    c_t6.button("ğŸ–¨ï¸ CertidÃ£o", use_container_width=True, on_click=toggle_view, args=("certidao",))
    c_t7.button("ğŸ’¡ SugestÃ£o", use_container_width=True, on_click=toggle_view, args=("sugestao",))

    if st.session_state.active_view == "checklist":
        with st.container(border=True):
            st.header("Gerador de Checklist"); data_eproc = st.date_input("Data:", value=get_brazil_time().date()); camara_eproc = st.text_input("CÃ¢mara:")
            if st.button("Gerar HTML"): send_to_chat("sessao", f"Consultor {st.session_state.consultor_selectbox} acompanhando sessÃ£o {camara_eproc}"); st.success("Registrado no chat!")
            if st.button("âŒ Cancelar"): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == "chamados":
        with st.container(border=True):
            st.header('ğŸ†˜ Chamados (PadrÃ£o / Jira)')
            # Etapas simplificadas - Direto para o texto
            st.text_area('Texto do chamado:', height=240, key='chamado_textarea')
            c1, c2 = st.columns(2)
            with c1:
                if st.button('Enviar', type='primary', use_container_width=True): # BotÃ£o renomeado
                    if handle_chamado_submission(): st.success('Enviado!'); st.session_state.active_view = None; st.rerun()
                    else: st.error('Erro ao enviar.')
            with c2:
                    if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == "atendimentos":
        with st.container(border=True):
            st.header('ğŸ“ Registro de Atendimentos')
            at_data = st.date_input('Data:', value=get_brazil_time().date())
            at_usuario = st.selectbox('UsuÃ¡rio:', REG_USUARIO_OPCOES); at_setor = st.text_input('Setor:'); at_sys = st.selectbox('Sistema:', REG_SISTEMA_OPCOES)
            at_desc = st.text_input('DescriÃ§Ã£o:'); at_canal = st.selectbox('Canal:', REG_CANAL_OPCOES); at_res = st.selectbox('Desfecho:', REG_DESFECHO_OPCOES); at_jira = st.text_input('Jira:')
            if st.button('Enviar', type='primary', use_container_width=True):
                if send_atendimento_to_chat(st.session_state.consultor_selectbox, at_data, at_usuario, at_setor, at_sys, at_desc, at_canal, at_res, at_jira):
                    st.success('Enviado!'); st.session_state.active_view = None; st.rerun()
                else: st.error('Erro.')
            if st.button('âŒ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == "hextras":
        with st.container(border=True):
            st.header("â° Horas Extras"); d_ex = st.date_input("Data:"); h_in = st.time_input("InÃ­cio:"); t_ex = st.text_input("Tempo Total:"); mot = st.text_input("Motivo:")
            if st.button("Registrar"): 
                if send_horas_extras_to_chat(st.session_state.consultor_selectbox, d_ex, h_in, t_ex, mot): st.success("Registrado!"); st.session_state.active_view = None; st.rerun()
            if st.button("âŒ Cancelar"): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == "erro_novidade":
        with st.container(border=True):
            st.header("ğŸ› Erro/Novidade"); tit = st.text_input("TÃ­tulo:"); obj = st.text_area("Objetivo:"); rel = st.text_area("Relato:"); res = st.text_area("Resultado:")
            if st.button("Enviar"): 
                if handle_erro_novidade_submission(st.session_state.consultor_selectbox, tit, obj, rel, res): st.success("Enviado!"); st.session_state.active_view = None; st.rerun()
            if st.button("âŒ Cancelar"): st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == "certidao":
        with st.container(border=True):
            st.header("ğŸ–¨ï¸ Registro de CertidÃ£o (2026)")
            # CorreÃ§Ã£o 1: Formato data DD/MM/YYYY
            c_data = st.date_input("Data do Evento:", value=get_brazil_time().date(), format="DD/MM/YYYY")
            tipo_cert = st.selectbox("Tipo:", ["FÃ­sica", "EletrÃ´nica", "Geral"])
            c_cons = st.session_state.consultor_selectbox
            
            # Campos comuns
            c_hora = "" # Mantido localmente para o DOCX se necessÃ¡rio, mas nÃ£o vai pro DB como coluna separada
            c_motivo = st.text_area("Motivo/Detalhes:", height=100)
            
            # Condicionais
            if tipo_cert == "Geral": 
                c_hora = st.text_input("HorÃ¡rio/PerÃ­odo (Ex: 13h Ã s 15h):")
                # Se for geral, nÃ£o tem processo especÃ­fico obrigatÃ³rio, mas vamos manter vazio
                c_proc = ""; c_chamado = ""; c_nome_parte = ""; c_peticao = ""
                # Concatena hora no motivo para salvar no banco
                if c_hora: c_motivo = f"{c_motivo} - PerÃ­odo: {c_hora}"
            else: 
                c1, c2 = st.columns(2)
                c_proc = c1.text_input("Processo (Com pontuaÃ§Ã£o):")
                c_chamado = c2.text_input("Incidente/Chamado:")
                
                c3, c4 = st.columns(2)
                c_nome_parte = c3.text_input("Nome da Parte/Advogado:")
                c_peticao = c4.selectbox("Tipo de PetiÃ§Ã£o:", ["Inicial", "Recursal", "IntermediÃ¡ria", "Outros"])
            
            c1, c2 = st.columns(2)
            with c1:
                # Gerar Word usa os dados locais para preencher o modelo
                if st.button("ğŸ“„ Gerar Word", use_container_width=True): 
                    st.session_state.word_buffer = gerar_docx_certidao_internal(tipo_cert, c_proc, c_data.strftime("%d/%m/%Y"), c_cons, c_motivo, c_chamado, c_hora, c_nome_parte)
                if st.session_state.word_buffer: 
                    st.download_button("â¬‡ï¸ Baixar", st.session_state.word_buffer, file_name="certidao.docx")
            with c2:
                if st.button("ğŸ’¾ Salvar e Notificar", type="primary", use_container_width=True):
                    # ValidaÃ§Ã£o de Duplicidade corrigida (pela coluna PROCESSO)
                    if verificar_duplicidade_certidao(tipo_cert, c_proc, c_data): 
                        st.session_state.aviso_duplicidade = True
                    else:
                        # Payload corrigido para o banco 'certidoes_registro'
                        payload = {
                            "tipo": tipo_cert, 
                            "data": c_data.isoformat(), 
                            "consultor": c_cons, 
                            "incidente": c_chamado, # Mapeia chamado -> incidente
                            "processo": c_proc, 
                            "motivo": c_motivo,
                            "nome_parte": c_nome_parte,
                            "peticao": c_peticao
                        }
                        if salvar_certidao_db(payload):
                            msg_cert = f"ğŸ–¨ï¸ **Nova CertidÃ£o Registrada**\nğŸ‘¤ **Autor:** {c_cons}\nğŸ“… **Data:** {c_data.strftime('%d/%m/%Y')}\nğŸ“„ **Tipo:** {tipo_cert}\nğŸ“‚ **Proc:** {c_proc}"
                            try: send_to_chat("certidao", msg_cert)
                            except Exception as e: st.error(f"Erro Webhook: {e}")
                            st.success("Salvo!"); time.sleep(1); st.session_state.active_view = None; st.session_state.word_buffer = None; st.rerun()
                        else: st.error("Erro ao salvar no banco.")
            
            if st.button("âŒ Cancelar"): st.session_state.active_view = None; st.rerun()
            if st.session_state.get('aviso_duplicidade'): st.error("âš ï¸ Este processo jÃ¡ possui registro de certidÃ£o!"); st.button("Ok, entendi", on_click=st.rerun)

    if st.session_state.active_view == "sugestao":
        with st.container(border=True):
            st.header("ğŸ’¡ Enviar SugestÃ£o")
            sug_txt = st.text_area("Sua ideia ou melhoria:")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Enviar SugestÃ£o", type="primary", use_container_width=True):
                    if handle_sugestao_submission(st.session_state.consultor_selectbox, sug_txt):
                        st.success("Enviado com sucesso!")
                        st.session_state.active_view = None; st.rerun()
                    else: st.error("Erro ao enviar.")
            with c2:
                if st.button("Cancelar", use_container_width=True): st.session_state.active_view = None; st.rerun()
    
    # --- GRÃFICO OPERACIONAL (DENTRO DA COLUNA ESQUERDA, NO FIM) ---
    st.markdown("---")
    st.subheader("ğŸ“Š Resumo Operacional")
    
    df_chart, gerado_em = carregar_dados_grafico()
    
    if df_chart is not None:
        try:
            df_long = df_chart.melt(id_vars=['relatorio'], value_vars=['Eproc', 'Legados'], var_name='Sistema', value_name='Qtd')
            base = alt.Chart(df_long).encode(
                x=alt.X('relatorio', title=None, axis=alt.Axis(labels=True, labelAngle=0)),
                y=alt.Y('Qtd', title='Quantidade'),
                color=alt.Color('Sistema', legend=alt.Legend(title="Sistema")),
                xOffset='Sistema'
            )
            bars = base.mark_bar()
            text = base.mark_text(dy=-5, color='black').encode(text='Qtd')
            final_chart = (bars + text).properties(height=300)
            st.altair_chart(final_chart, use_container_width=True)
            st.caption(f"Dados do dia: {gerado_em} (AtualizaÃ§Ã£o diÃ¡ria)")
            st.markdown("### Dados Detalhados")
            st.dataframe(df_chart, use_container_width=True)
        except Exception as e: st.error(f"Erro grÃ¡fico: {e}")
    else: st.info("Sem dados de resumo disponÃ­veis.")
