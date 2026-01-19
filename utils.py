import streamlit as st
import requests
# import threading  <-- REMOVIDO: Não vamos mais usar threads
import base64
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

# --- CONSTANTES ---
CONSULTORES = sorted([
    "Alex Paulo", "Dirceu Gonçalves", "Douglas De Souza", "Farley Leandro", "Gleis Da Silva", 
    "Hugo Leonardo", "Igor Dayrell", "Jerry Marcos", "Jonatas Gomes", "Leandro Victor", 
    "Luiz Henrique", "Marcelo Dos Santos", "Marina Silva", "Marina Torres", "Vanessa Ligiane"
])

def get_secret(section, key):
    try: return st.secrets[section][key]
    except: return ""

def get_brazil_time():
    return datetime.utcnow() - timedelta(hours=3)

def _send_webhook(url, payload):
    """
    Envia o webhook de forma SÍNCRONA (bloqueante) mas com timeout curto.
    Isso garante que o Streamlit Cloud não mate o processo antes do envio.
    """
    if not url: return
    try:
        # Timeout de 3 segundos para não travar o app se o Google cair
        requests.post(url, json=payload, headers={'Content-Type': 'application/json'}, timeout=3)
    except Exception as e:
        print(f"Erro ao enviar webhook: {e}")

def send_to_chat(webhook_key, text_msg):
    url = get_secret("chat", webhook_key)
    if not url: return False
    
    # [CORREÇÃO] Chamada direta sem Threading
    _send_webhook(url, {"text": text_msg})
    
    return True

def gerar_docx_certidao(tipo_certidao, num_processo, data_indisponibilidade_input, num_chamado, motivo_pedido):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    head = document.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tj = head.add_run("TRIBUNAL DE JUSTIÇA DO ESTADO DE MINAS GERAIS\n")
    run_tj.bold = True
    head.add_run("Rua Ouro Preto, Nº 1564 - Bairro Santo Agostinho - CEP 30170-041 - Belo Horizonte - MG - www.tjmg.jus.br\n")
    head.add_run("Andar: 3º 3º e 4º PV\n\n")

    num_parecer = int(datetime.now().strftime("%H%M"))
    ano_atual = datetime.now().year
    titulo = document.add_paragraph(f"Parecer Técnico GEJUD/DIRTEC/TJMG nº {num_parecer}/{ano_atual}.")
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    titulo.runs[0].bold = True

    document.add_paragraph(f"Assunto: Notifica erro no JPe ao peticionar.")
    document.add_paragraph(f"Motivo: {motivo_pedido}")
    
    if num_processo:
        document.add_paragraph(f"Processo: {num_processo}")
    if num_chamado:
        document.add_paragraph(f"Chamado: {num_chamado}")

    document.add_paragraph("\nDocumento gerado automaticamente pelo Sistema de Controle de Bastão.")
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

@st.cache_data
def get_img_as_base64(file_path):
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except: return None
